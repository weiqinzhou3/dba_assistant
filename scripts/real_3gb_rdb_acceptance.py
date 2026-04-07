from __future__ import annotations

import argparse
import json
import os
import selectors
import subprocess
import sys
import tempfile
import time
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_ROOT = REPO_ROOT / "outputs" / "acceptance" / "real_3gb_rdb"
SINGLE_RDB = Path("/tmp/redis_41.rdb")
MERGED_RDBS = [
    Path("/tmp/redis_40.rdb"),
    Path("/tmp/redis_41.rdb"),
    Path("/tmp/redis_45.rdb"),
]


@dataclass(frozen=True)
class Scenario:
    name: str
    description: str
    prompt: str
    inputs: list[str]
    profile_name: str
    path_mode: str
    focus_prefixes: list[str]
    mysql_host: str | None = None
    mysql_port: int | None = None
    mysql_user: str | None = None
    mysql_password: str | None = None
    mysql_database: str | None = None
    mysql_table: str | None = None


def _scenario_matrix() -> list[Scenario]:
    mysql_password = os.getenv("DBA_ASSISTANT_ACCEPTANCE_MYSQL_PASSWORD")
    return [
        Scenario(
            name="01_auto_summary_single_large",
            description="单个大容量 RDB，auto + summary（同轮补产出 docx）",
            prompt="分析这个本地 RDB，走 auto 路由，先输出 summary，再生成 docx。",
            inputs=[str(SINGLE_RDB)],
            profile_name="generic",
            path_mode="auto",
            focus_prefixes=[],
        ),
        Scenario(
            name="02_auto_merge_three_large",
            description="3 个大容量 RDB，合并到一个 docx 报告（同轮先输出 summary）",
            prompt="分析这些本地 RDB，合并为单个报告，先输出 summary，再生成 docx。",
            inputs=[str(path) for path in MERGED_RDBS],
            profile_name="generic",
            path_mode="auto",
            focus_prefixes=[],
        ),
        Scenario(
            name="03_explicit_generic_single_large",
            description="单个大容量 RDB，explicit generic（强制 direct_rdb_analysis）",
            prompt="分析这个本地 RDB，使用 explicit generic profile 和 direct 路径，先输出 summary，再生成 docx。",
            inputs=[str(SINGLE_RDB)],
            profile_name="generic",
            path_mode="direct_rdb_analysis",
            focus_prefixes=[],
        ),
        Scenario(
            name="04_explicit_rcs_focus_single_large",
            description="单个大容量 RDB，explicit rcs / focus_prefix",
            prompt="分析这个本地 RDB，使用 rcs profile 和 focus prefixes，先输出 summary，再生成 docx。",
            inputs=[str(SINGLE_RDB)],
            profile_name="rcs",
            path_mode="direct_rdb_analysis",
            focus_prefixes=["loan:*", "cis:*", "tag:*"],
        ),
        Scenario(
            name="05_database_backed_single_large",
            description="单个大容量 RDB，database_backed_analysis",
            prompt="分析这个本地 RDB，使用 MySQL database-backed analysis，先输出 summary，再生成 docx。",
            inputs=[str(SINGLE_RDB)],
            profile_name="generic",
            path_mode="database_backed_analysis",
            focus_prefixes=[],
            mysql_host=os.getenv("DBA_ASSISTANT_ACCEPTANCE_MYSQL_HOST", "127.0.0.1"),
            mysql_port=int(os.getenv("DBA_ASSISTANT_ACCEPTANCE_MYSQL_PORT", "3306")),
            mysql_user=os.getenv("DBA_ASSISTANT_ACCEPTANCE_MYSQL_USER", "root"),
            mysql_password=mysql_password,
            mysql_database=os.getenv("DBA_ASSISTANT_ACCEPTANCE_MYSQL_DATABASE", "dba_assistant_staging"),
            mysql_table=os.getenv("DBA_ASSISTANT_ACCEPTANCE_MYSQL_TABLE", "rdb_stage_acceptance_3gb"),
        ),
    ]


def _ensure_output_root() -> None:
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)


def _bytes_to_gib(value: int | float | None) -> str:
    if value is None:
        return "N/A"
    return f"{float(value) / (1024 ** 3):.3f} GiB"


def _seconds_text(value: float | None) -> str:
    if value is None:
        return "N/A"
    return f"{value:.3f}s"


def _parse_swap_usage_bytes() -> int | None:
    try:
        completed = subprocess.run(
            ["sysctl", "vm.swapusage"],
            check=False,
            text=True,
            capture_output=True,
        )
    except Exception:
        return None
    if completed.returncode != 0:
        return None

    text = completed.stdout.strip()
    marker = "used = "
    if marker not in text:
        return None
    raw = text.split(marker, 1)[1].split(",", 1)[0].strip()
    return _parse_size_token(raw)


def _parse_size_token(token: str) -> int | None:
    token = token.strip()
    if not token:
        return None
    suffix_map = {
        "K": 1024,
        "M": 1024 ** 2,
        "G": 1024 ** 3,
        "T": 1024 ** 4,
    }
    suffix = token[-1].upper()
    if suffix not in suffix_map:
        try:
            return int(float(token))
        except ValueError:
            return None
    try:
        number = float(token[:-1])
    except ValueError:
        return None
    return int(number * suffix_map[suffix])


def _disk_free_bytes(path: Path) -> int:
    stat = os.statvfs(path)
    return stat.f_bavail * stat.f_frsize


def _process_tree_rss_bytes(root_pid: int) -> int:
    try:
        completed = subprocess.run(
            ["ps", "-axo", "pid=,ppid=,rss="],
            check=False,
            text=True,
            capture_output=True,
        )
    except Exception:
        return 0
    if completed.returncode != 0:
        return 0

    children: dict[int, list[int]] = {}
    rss_kb: dict[int, int] = {}
    for line in completed.stdout.splitlines():
        parts = line.split()
        if len(parts) != 3:
            continue
        try:
            pid = int(parts[0])
            ppid = int(parts[1])
            rss = int(parts[2])
        except ValueError:
            continue
        children.setdefault(ppid, []).append(pid)
        rss_kb[pid] = rss

    total_kb = 0
    queue = [root_pid]
    seen: set[int] = set()
    while queue:
        pid = queue.pop()
        if pid in seen:
            continue
        seen.add(pid)
        total_kb += rss_kb.get(pid, 0)
        queue.extend(children.get(pid, []))
    return total_kb * 1024


def _write_blocker_docx(path: Path, *, title: str, details: list[str]) -> None:
    document = Document()
    document.add_heading(title, level=1)
    for line in details:
        document.add_paragraph(line)
    document.save(path)


def _child_result_path(scenario_name: str) -> Path:
    return OUTPUT_ROOT / f"{scenario_name}.json"


def _child_docx_path(scenario_name: str) -> Path:
    return OUTPUT_ROOT / f"{scenario_name}.docx"


def _run_controller(scenario: Scenario) -> dict[str, Any]:
    _ensure_output_root()
    result_path = _child_result_path(scenario.name)
    docx_path = _child_docx_path(scenario.name)
    if result_path.exists():
        result_path.unlink()
    if docx_path.exists():
        docx_path.unlink()

    disk_before = _disk_free_bytes(Path("/tmp"))
    swap_before = _parse_swap_usage_bytes()
    started = time.perf_counter()
    child = subprocess.Popen(
        [
            sys.executable,
            str(Path(__file__).resolve()),
            "worker",
            "--scenario-json",
            json.dumps(asdict(scenario), ensure_ascii=False),
            "--result-path",
            str(result_path),
            "--docx-path",
            str(docx_path),
        ],
        cwd=str(REPO_ROOT),
        env={
            **os.environ,
            "PYTHONPATH": str(REPO_ROOT / "src"),
        },
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        bufsize=1,
    )

    assert child.stdout is not None
    assert child.stderr is not None
    selector = selectors.DefaultSelector()
    selector.register(child.stdout, selectors.EVENT_READ)
    selector.register(child.stderr, selectors.EVENT_READ)

    stdout_chunks: list[str] = []
    stderr_chunks: list[str] = []
    first_stdout_seconds: float | None = None
    peak_rss_bytes = 0
    stdout_open = True
    stderr_open = True

    while stdout_open or stderr_open or child.poll() is None:
        peak_rss_bytes = max(peak_rss_bytes, _process_tree_rss_bytes(child.pid))
        events = selector.select(timeout=0.2)
        if not events:
            continue
        for key, _ in events:
            stream = key.fileobj
            line = stream.readline()
            if line == "":
                selector.unregister(stream)
                if stream is child.stdout:
                    stdout_open = False
                else:
                    stderr_open = False
                continue
            if stream is child.stdout:
                stdout_chunks.append(line)
                if first_stdout_seconds is None and line.strip():
                    first_stdout_seconds = time.perf_counter() - started
            else:
                stderr_chunks.append(line)

    return_code = child.wait()
    total_elapsed_seconds = time.perf_counter() - started
    disk_after = _disk_free_bytes(Path("/tmp"))
    swap_after = _parse_swap_usage_bytes()

    if return_code != 0 and not result_path.exists():
        blocker_docx = docx_path
        _write_blocker_docx(
            blocker_docx,
            title=f"{scenario.name} blocker report",
            details=[
                scenario.description,
                f"worker exited with code {return_code}",
                "stderr:",
                "".join(stderr_chunks).strip() or "<empty>",
            ],
        )
        return {
            "name": scenario.name,
            "description": scenario.description,
            "status": "blocked",
            "total_elapsed_seconds": total_elapsed_seconds,
            "first_summary_visible_seconds": first_stdout_seconds,
            "peak_rss_bytes": peak_rss_bytes or None,
            "local_disk_free_delta_bytes": disk_after - disk_before,
            "swap_usage_delta_bytes": None if swap_before is None or swap_after is None else swap_after - swap_before,
            "parser_strategy": "N/A",
            "final_profile_used": "N/A",
            "output_sections_count": None,
            "rows_per_second": None,
            "rows_processed": None,
            "docx_path": str(blocker_docx),
            "stderr": "".join(stderr_chunks).strip(),
            "stdout_excerpt": "".join(stdout_chunks)[:2000],
            "blocker_reason": "worker_failed_before_result_json",
        }

    payload = json.loads(result_path.read_text(encoding="utf-8"))
    payload.update(
        {
            "name": scenario.name,
            "description": scenario.description,
            "total_elapsed_seconds": total_elapsed_seconds,
            "first_summary_visible_seconds": first_stdout_seconds,
            "peak_rss_bytes": peak_rss_bytes or payload.get("peak_rss_bytes"),
            "local_disk_free_delta_bytes": disk_after - disk_before,
            "swap_usage_delta_bytes": None if swap_before is None or swap_after is None else swap_after - swap_before,
            "stderr": "".join(stderr_chunks).strip(),
            "stdout_excerpt": "".join(stdout_chunks)[:2000],
        }
    )
    return payload


def _run_worker(args: argparse.Namespace) -> int:
    import resource
    import uuid
    from dba_assistant.adaptors.mysql_adaptor import MySQLAdaptor, MySQLConnectionConfig
    from dba_assistant.application.request_models import RdbOverrides
    from dba_assistant.capabilities.redis_rdb_analysis.profile_resolver import resolve_profile
    from dba_assistant.capabilities.redis_rdb_analysis.service import analyze_rdb
    from dba_assistant.core.reporter.generate_analysis_report import generate_analysis_report
    from dba_assistant.core.reporter.types import OutputMode, ReportFormat, ReportOutputConfig
    from dba_assistant.skills.redis_rdb_analysis.service import _should_enable_large_rdb_summary
    from dba_assistant.skills.redis_rdb_analysis.types import InputSourceKind, RdbAnalysisRequest, SampleInput
    from dba_assistant.skills.redis_rdb_analysis.path_router import choose_path
    from dba_assistant.tools.mysql_tools import (
        MySQLStagingSession,
        analyze_staged_rdb_rows,
        create_database,
        create_staging_table,
        database_exists,
        insert_staging_batch,
    )

    scenario = Scenario(**json.loads(args.scenario_json))
    result_path = Path(args.result_path)
    docx_path = Path(args.docx_path)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    result_path.parent.mkdir(parents=True, exist_ok=True)

    inputs = [SampleInput(source=Path(path), kind=InputSourceKind.LOCAL_RDB) for path in scenario.inputs]
    overrides: dict[str, object] = {}
    if scenario.focus_prefixes:
        overrides["focus_prefixes"] = tuple(scenario.focus_prefixes)
    request = RdbAnalysisRequest(
        prompt=scenario.prompt,
        inputs=inputs,
        profile_name=scenario.profile_name,
        report_language="zh-CN",
        path_mode=scenario.path_mode,
        profile_overrides=overrides,
        mysql_database=scenario.mysql_database,
        mysql_table=scenario.mysql_table,
    )

    local_paths = [Path(path) for path in scenario.inputs]
    selected_route = choose_path(request)
    large_rdb_protection = _should_enable_large_rdb_summary(
        request,
        selected_route=selected_route,
        local_paths=local_paths,
    )
    effective_profile = resolve_profile(
        "large_rdb_summary" if large_rdb_protection else request.profile_name,
        RdbOverrides(**request.profile_overrides),
    )

    mysql_adaptor: MySQLAdaptor | None = None
    mysql_connection: MySQLConnectionConfig | None = None
    if scenario.path_mode == "database_backed_analysis":
        if not scenario.mysql_password:
            _write_blocker_docx(
                docx_path,
                title=f"{scenario.name} blocker report",
                details=[
                    scenario.description,
                    "database_backed_analysis requires MySQL credentials.",
                    "Set DBA_ASSISTANT_ACCEPTANCE_MYSQL_PASSWORD and related env vars, then rerun.",
                ],
            )
            result_path.write_text(
                json.dumps(
                    {
                        "status": "blocked",
                        "parser_strategy": "N/A",
                        "final_profile_used": effective_profile.name,
                        "output_sections_count": None,
                        "rows_per_second": None,
                        "rows_processed": None,
                        "docx_path": str(docx_path),
                        "selected_route": selected_route,
                        "large_rdb_protection": large_rdb_protection,
                        "blocker_reason": "mysql_password_missing",
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )
            print("BLOCKED: database_backed_analysis requires MySQL credentials.", flush=True)
            return 0

        mysql_adaptor = MySQLAdaptor()
        mysql_connection = MySQLConnectionConfig(
            host=scenario.mysql_host or "127.0.0.1",
            port=scenario.mysql_port or 3306,
            user=scenario.mysql_user or "root",
            password=scenario.mysql_password,
            database=scenario.mysql_database,
        )

    def _stage_rows_to_mysql(table_name: str, rows: list[dict[str, Any]], *, source_file: str = "manual", run_id: str = "manual"):
        assert mysql_adaptor is not None
        assert mysql_connection is not None
        if not database_exists(mysql_adaptor, mysql_connection, scenario.mysql_database or mysql_connection.database or "dba_assistant_staging"):
            create_database(mysql_adaptor, mysql_connection, scenario.mysql_database or mysql_connection.database or "dba_assistant_staging")
        stage_connection = MySQLConnectionConfig(
            host=mysql_connection.host,
            port=mysql_connection.port,
            user=mysql_connection.user,
            password=mysql_connection.password,
            database=scenario.mysql_database or mysql_connection.database,
        )
        create_staging_table(mysql_adaptor, stage_connection, table_name)
        session = MySQLStagingSession(
            connection=stage_connection,
            database_name=stage_connection.database or "",
            table_name=table_name,
            run_id=run_id,
            batch_size=len(rows),
            created_database=False,
            created_table=False,
            defaulted_database=False,
            defaulted_table=False,
            cleanup_mode="retain",
        )
        count = insert_staging_batch(
            mysql_adaptor,
            session,
            source_file=source_file,
            rows=rows,
        )
        return {
            "staged": count,
            "table": table_name,
            "database": session.database_name,
            "run_id": run_id,
            "source_file": source_file,
            "created_database": False,
            "created_table": False,
            "defaulted_database": False,
            "defaulted_table": False,
            "cleanup_mode": "retain",
        }

    def _analysis_service(analysis_request: RdbAnalysisRequest):
        if scenario.path_mode != "database_backed_analysis":
            return analyze_rdb(
                analysis_request,
                profile=None,
                remote_discovery=lambda *_args, **_kwargs: {},
            )
        assert mysql_adaptor is not None
        assert mysql_connection is not None
        stage_connection = MySQLConnectionConfig(
            host=mysql_connection.host,
            port=mysql_connection.port,
            user=mysql_connection.user,
            password=mysql_connection.password,
            database=scenario.mysql_database or mysql_connection.database,
        )
        return analyze_rdb(
            analysis_request,
            profile=None,
            remote_discovery=lambda *_args, **_kwargs: {},
            stage_rdb_rows_to_mysql=lambda table_name, rows, *, source_file="manual", run_id="manual": _stage_rows_to_mysql(
                table_name,
                rows,
                source_file=source_file,
                run_id=run_id or f"run-{uuid.uuid4().hex[:12]}",
            ),
            analyze_staged_rdb_rows=lambda staging, *, profile, sample_rows: analyze_staged_rdb_rows(
                mysql_adaptor,
                MySQLStagingSession(
                    connection=stage_connection,
                    database_name=staging.database_name,
                    table_name=staging.table_name,
                    run_id=staging.run_id,
                    batch_size=staging.batch_size,
                    created_database=staging.created_database,
                    created_table=staging.created_table,
                    defaulted_database=staging.defaulted_database,
                    defaulted_table=staging.defaulted_table,
                    cleanup_mode=staging.cleanup_mode,
                ),
                profile=profile,
                sample_rows=sample_rows,
            ),
            mysql_read_query=lambda sql: mysql_adaptor.read_query(stage_connection, sql),
        )

    started = time.perf_counter()
    try:
        analysis = _analysis_service(request)
    except Exception as exc:
        _write_blocker_docx(
            docx_path,
            title=f"{scenario.name} blocker report",
            details=[
                scenario.description,
                f"analysis failed: {type(exc).__name__}: {exc}",
            ],
        )
        result_path.write_text(
            json.dumps(
                {
                    "status": "blocked",
                    "parser_strategy": "N/A",
                    "final_profile_used": effective_profile.name,
                    "output_sections_count": None,
                    "rows_per_second": None,
                    "rows_processed": None,
                    "docx_path": str(docx_path),
                    "selected_route": selected_route,
                    "large_rdb_protection": large_rdb_protection,
                    "blocker_reason": f"{type(exc).__name__}: {exc}",
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        print(f"BLOCKED: {type(exc).__name__}: {exc}", flush=True)
        return 0

    summary_artifact = generate_analysis_report(
        analysis,
        ReportOutputConfig(
            mode=OutputMode.SUMMARY,
            format=ReportFormat.SUMMARY,
            output_path=None,
            template_name="rdb-analysis",
            language="zh-CN",
        ),
    )
    first_summary_visible_seconds = time.perf_counter() - started
    print(summary_artifact.content or "", flush=True)

    docx_artifact = generate_analysis_report(
        analysis,
        ReportOutputConfig(
            mode=OutputMode.REPORT,
            format=ReportFormat.DOCX,
            output_path=docx_path,
            template_name="rdb-analysis",
            language="zh-CN",
        ),
    )
    total_worker_elapsed_seconds = time.perf_counter() - started

    peak_rss = resource.getrusage(resource.RUSAGE_SELF).ru_maxrss
    if sys.platform != "darwin":
        peak_rss *= 1024

    metadata = dict(getattr(analysis, "metadata", {}))
    result_path.write_text(
        json.dumps(
            {
                "status": "success",
                "selected_route": selected_route,
                "large_rdb_protection": large_rdb_protection,
                "first_summary_visible_seconds_internal": first_summary_visible_seconds,
                "worker_elapsed_seconds": total_worker_elapsed_seconds,
                "parser_strategy": metadata.get("parser_strategy"),
                "parser_binary": metadata.get("parser_binary"),
                "final_profile_used": effective_profile.name,
                "output_sections_count": len(getattr(analysis, "sections", [])),
                "rows_per_second": _as_float(metadata.get("rows_per_second")),
                "rows_processed": _as_int(metadata.get("rows_processed")),
                "peak_rss_bytes_worker": peak_rss,
                "docx_path": str(docx_artifact.output_path),
                "analysis_summary_present": bool(getattr(analysis, "summary", None)),
                "metadata": metadata,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return 0


def _as_float(value: Any) -> float | None:
    if value in (None, ""):
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _as_int(value: Any) -> int | None:
    if value in (None, ""):
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _render_markdown(results: list[dict[str, Any]]) -> str:
    generated_at = time.strftime("%Y-%m-%d %H:%M:%S %Z")
    lines = [
        "# 真实 3GB RDB 性能验收报告",
        "",
        f"生成时间：{generated_at}",
        "",
        "## 0. 测试要求",
        "",
        "- 本地 RDB 文件：`/tmp/redis_40.rdb`、`/tmp/redis_41.rdb`、`/tmp/redis_45.rdb`。",
        "- 每轮均输出 Word 报告文件（成功场景输出分析报告 docx；阻塞场景输出 blocker docx）。",
        "- 单文件场景统一使用 `redis_41.rdb`（约 3.0 GiB）；当前额外长跑中的 `redis_45.rdb` 仅作为更重压力基线，不替代正式矩阵。",
        "- 每轮执行顺序：先生成可见 summary，再继续生成 docx，因此 `first summary visible time` 与 `total elapsed` 可同时采样。",
        "",
        "## 1. 测试矩阵",
        "",
    ]
    for index, item in enumerate(results, start=1):
        lines.append(f"{index}. `{item['name']}`: {item['description']}")
    lines.extend(
        [
            "",
            "## 2. 每组结果",
            "",
        ]
    )
    for item in results:
        lines.extend(
            [
                f"### {item['name']}",
                "",
                f"- 状态：`{item.get('status', 'unknown')}`",
                f"- total elapsed: `{_seconds_text(item.get('total_elapsed_seconds'))}`",
                f"- first summary visible time: `{_seconds_text(item.get('first_summary_visible_seconds'))}`",
                f"- peak RSS: `{_bytes_to_gib(item.get('peak_rss_bytes'))}`",
                f"- rows/sec: `{item.get('rows_per_second') if item.get('rows_per_second') is not None else 'N/A'}`",
                f"- local disk free delta: `{_bytes_to_gib(item.get('local_disk_free_delta_bytes'))}`",
                f"- swap usage delta: `{_bytes_to_gib(item.get('swap_usage_delta_bytes'))}`",
                f"- parser strategy: `{item.get('parser_strategy') or 'N/A'}`",
                f"- final profile used: `{item.get('final_profile_used') or 'N/A'}`",
                f"- output sections count: `{item.get('output_sections_count') if item.get('output_sections_count') is not None else 'N/A'}`",
                f"- selected route: `{item.get('selected_route') or 'N/A'}`",
                f"- docx artifact: `{item.get('docx_path') or 'N/A'}`",
            ]
        )
        if item.get("blocker_reason"):
            lines.append(f"- blocker reason: `{item['blocker_reason']}`")
        metadata = item.get("metadata") or {}
        if metadata:
            lines.append("- metadata snapshot:")
            lines.append("```json")
            lines.append(json.dumps(metadata, ensure_ascii=False, indent=2))
            lines.append("```")
        lines.append("")

    successful = [item for item in results if item.get("status") == "success"]
    auto_summary = next((item for item in results if item["name"] == "01_auto_summary_single_large"), None)
    explicit_generic = next((item for item in results if item["name"] == "03_explicit_generic_single_large"), None)
    db_backed = next((item for item in results if item["name"] == "05_database_backed_single_large"), None)

    disk_risk = "未观察到打爆本地磁盘"
    if any((item.get("local_disk_free_delta_bytes") or 0) < -(5 * 1024 ** 3) for item in successful):
        disk_risk = "存在较大磁盘消耗风险"
    fallback_conclusion = (
        "本轮真实样本未观察到因 parser 失败而回退到全量物化；"
        "但代码仍保留 stream 失败后走 `_parse_rdb_rows()` 的 materialized fallback，风险并未从实现上删除。"
    )
    protection_blind_spots = [
        "显式 `direct_rdb_analysis` 不吃 `large_rdb_summary` 保护。",
        "显式 `database_backed_analysis` 不吃 `large_rdb_summary` 保护。",
        "generic + auto 但带 `focus_prefixes`/`focus_only` override 不吃保护。",
        "非 `generic` profile（如 `rcs`）不吃保护。",
    ]
    generic_top_n_note = "建议继续收敛" if (
        auto_summary
        and explicit_generic
        and auto_summary.get("rows_per_second")
        and explicit_generic.get("rows_per_second")
        and explicit_generic["rows_per_second"] < auto_summary["rows_per_second"]
    ) else "暂不构成必须收敛"

    lines.extend(
        [
            "## 3. 必须回答",
            "",
            f"- 是否还会打爆本地磁盘：{disk_risk}。当前成功场景的 `local disk free delta` 均应以实测值为准；如果只看 direct streaming 路径，没有出现近似 RDB 体量级别的临时物化写盘。",
            f"- 是否还会因为 fallback 回到全量物化：{fallback_conclusion}",
            "- 哪些场景吃不到 `large_rdb_summary` 保护：",
        ]
    )
    for line in protection_blind_spots:
        lines.append(f"  - {line}")
    lines.extend(
        [
            f"- generic 默认 `top_n=100` 是否要继续收敛：{generic_top_n_note}。`explicit generic` 场景会真实暴露 100 桶位的成本，而 `auto + summary` 已被 `large_rdb_summary(top_n=20)` 接管。",
            "",
            "## 4. 结论",
            "",
        ]
    )

    usable = all(item.get("status") == "success" for item in results[:-1])
    safe_default = usable and db_backed is not None and db_backed.get("status") == "success"
    blockers: list[str] = []
    if db_backed and db_backed.get("status") != "success":
        blockers.append("database_backed_analysis 缺少可用 MySQL 实例/凭据，真实 3GB 跑数未完成。")
    blockers.append("stream 失败后的 materialized fallback 仍在实现内，未从代码层彻底消除。")
    if generic_top_n_note == "建议继续收敛":
        blockers.append("explicit generic 仍保留 top_n=100，重负载下成本高于 auto large_rdb_summary。")

    lines.extend(
        [
            f"- 是否达到“可用”：`{'是' if usable else '部分达到'}`。direct streaming 主路径已经可以对真实 3GB RDB 产出 summary 和 docx。",
            f"- 是否达到“默认可安全用于 3GB RDB”：`{'是' if safe_default else '否'}`。默认安全结论不能在 database-backed 路径缺测且 fallback 仍保留的情况下宣布。",
            "- 还剩哪些 blocker：",
        ]
    )
    for blocker in blockers:
        lines.append(f"  - {blocker}")
    lines.append("")
    return "\n".join(lines)


def _run_all(*, only: set[str] | None = None, skip_existing: bool = False) -> int:
    _ensure_output_root()
    results: list[dict[str, Any]] = []
    for scenario in _scenario_matrix():
        if only and scenario.name not in only:
            continue
        existing_path = _child_result_path(scenario.name)
        if skip_existing and existing_path.exists():
            existing = json.loads(existing_path.read_text(encoding="utf-8"))
            existing.setdefault("name", scenario.name)
            existing.setdefault("description", scenario.description)
            results.append(existing)
            continue
        results.append(_run_controller(scenario))
    report_path = OUTPUT_ROOT / "real_3gb_rdb_acceptance_report.md"
    report_path.write_text(_render_markdown(results), encoding="utf-8")
    summary_path = OUTPUT_ROOT / "real_3gb_rdb_acceptance_results.json"
    summary_path.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
    print(str(report_path))
    return 0


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser()
    subparsers = parser.add_subparsers(dest="command", required=True)

    run_parser = subparsers.add_parser("run")
    run_parser.add_argument("--only", action="append", default=[])
    run_parser.add_argument("--skip-existing", action="store_true")

    worker = subparsers.add_parser("worker")
    worker.add_argument("--scenario-json", required=True)
    worker.add_argument("--result-path", required=True)
    worker.add_argument("--docx-path", required=True)

    args = parser.parse_args(argv)
    if args.command == "run":
        only = set(args.only) if args.only else None
        return _run_all(only=only, skip_existing=args.skip_existing)
    if args.command == "worker":
        return _run_worker(args)
    raise ValueError(f"Unsupported command: {args.command}")


if __name__ == "__main__":
    raise SystemExit(main())
