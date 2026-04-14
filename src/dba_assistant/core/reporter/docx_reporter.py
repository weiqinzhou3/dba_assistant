"""Professional DOCX reporter for repository-owned analysis reports."""

from __future__ import annotations

from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path
from types import ModuleType

from docx import Document

from dba_assistant.core.analyzer.types import AnalysisResult
from dba_assistant.core.reporter.docx_styles import (
    add_body_paragraph,
    add_cover_metadata,
    add_cover_subtitle,
    add_cover_title,
    add_heading,
    add_page_break,
    add_table_title,
    apply_document_theme,
    style_table,
)
from dba_assistant.core.reporter.localization import normalize_report_language
from dba_assistant.core.reporter.report_model import AnalysisReport, TableBlock, TextBlock, coerce_analysis_report
from dba_assistant.core.reporter.types import IReporter, ReportArtifact, ReportFormat, ReportOutputConfig


class DocxReporter(IReporter[AnalysisResult | AnalysisReport]):
    def __init__(self, repository_root: Path | None = None) -> None:
        self.repository_root = repository_root or Path(__file__).resolve().parents[4]

    def render(self, analysis: AnalysisResult | AnalysisReport, config: ReportOutputConfig) -> ReportArtifact:
        if config.output_path is None:
            raise ValueError("DocxReporter requires an output_path.")

        template_name = config.template_name or "rdb-analysis"
        template_module = self._load_module(
            self.repository_root / "templates" / "reports" / template_name / "template_spec.py",
            f"{template_name.replace('-', '_')}_template",
        )
        cover_module = self._load_module(
            self.repository_root / "templates" / "reports" / "shared" / "cover.py",
            "shared_cover",
        )
        disclaimer_module = self._load_module(
            self.repository_root / "templates" / "reports" / "shared" / "disclaimer.py",
            "shared_disclaimer",
        )
        table_style_module = self._load_module(
            self.repository_root / "templates" / "reports" / "shared" / "table_styles.py",
            "shared_table_styles",
        )

        language = self._resolve_language(analysis, config)
        template_text = self._resolve_text_map(template_module.TEMPLATE_TEXT, language)
        disclaimer_text = self._resolve_text_map(disclaimer_module.DISCLAIMER_TEXT, language)
        report = self._coerce_report(analysis, language)
        summary_text, sections = self._prepare_summary_and_sections(report)

        cover = cover_module.build_cover_spec(
            title=template_text["cover_title"],
            subtitle=template_text["cover_subtitle"],
        )

        document = Document()
        apply_document_theme(document, language=language)
        self._render_cover(document, cover, report.metadata, language)

        major_index = 0
        if summary_text:
            major_index += 1
            self._render_major_heading(document, major_index, str(template_text["summary_heading"]), language=language)
            add_body_paragraph(document, summary_text)

        major_index = self._render_sections(
            document,
            sections,
            table_style_module,
            language,
            start_major_index=major_index,
        )

        if template_module.TEMPLATE["include_disclaimer"]:
            major_index += 1
            self._render_major_heading(document, major_index, str(disclaimer_text["title"]), language=language)
            for paragraph in disclaimer_text["paragraphs"]:
                add_body_paragraph(document, str(paragraph))

        config.output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(config.output_path)
        return ReportArtifact(format=ReportFormat.DOCX, output_path=config.output_path, content=None)

    def _render_cover(self, document: Document, cover, metadata: dict[str, str], language: str) -> None:
        add_cover_title(document, cover.title)
        add_cover_subtitle(document, cover.subtitle)

        for key in cover.metadata_order:
            if key in metadata:
                add_cover_metadata(document, key, metadata[key], language=language)

        add_page_break(document)

    def _render_sections(
        self,
        document: Document,
        sections,
        table_style_module: ModuleType,
        language: str,
        *,
        start_major_index: int,
    ) -> int:
        major_index = start_major_index
        minor_index = 0
        sub_index = 0
        for section in sections:
            if section.level <= 1:
                major_index += 1
                minor_index = 0
                sub_index = 0
                self._render_major_heading(document, major_index, section.title, language=language)
            else:
                minor_index += 1
                sub_index = 0
                self._render_minor_heading(document, major_index, minor_index, section.title)
            for block in section.blocks:
                if isinstance(block, TextBlock):
                    add_body_paragraph(document, block.text)
                    continue
                if not isinstance(block, TableBlock):
                    raise TypeError(f"Unsupported block type: {type(block)!r}")
                if section.level >= 2 and block.title:
                    sub_index += 1
                    numbered_title = f"{major_index}.{minor_index}.{sub_index} {block.title}"
                else:
                    numbered_title = block.title
                add_table_title(document, numbered_title)
                docx_table = document.add_table(rows=1, cols=len(block.columns))
                for index, column in enumerate(block.columns):
                    docx_table.rows[0].cells[index].text = column
                for row in block.rows:
                    cells = docx_table.add_row().cells
                    for index, value in enumerate(row):
                        cells[index].text = value
                style_table(docx_table, language=language, table_style_module=table_style_module)
        return major_index

    def _render_major_heading(self, document: Document, index: int, title: str, *, language: str) -> None:
        prefix = _major_heading_prefix(index, language=language)
        add_heading(document, f"{prefix}{title}", level=1)

    def _render_minor_heading(self, document: Document, major: int, minor: int, title: str) -> None:
        add_heading(document, f"{major}.{minor} {title}", level=2)

    def _prepare_summary_and_sections(self, report: AnalysisReport) -> tuple[str | None, list]:
        summary_text = report.summary
        filtered_sections = [section for section in report.sections if section.id != "risk_summary"]

        executive_sections = [section for section in filtered_sections if section.id == "executive_summary" or section.title in {"执行摘要", "Executive Summary"}]
        if executive_sections:
            if summary_text is None:
                summary_text = "\n".join(
                    block.text
                    for section in executive_sections
                    for block in section.blocks
                    if isinstance(block, TextBlock) and block.text
                ) or None
            filtered_sections = [section for section in filtered_sections if section not in executive_sections]

        return summary_text, filtered_sections

    def _coerce_report(self, analysis: AnalysisResult | AnalysisReport, language: str) -> AnalysisReport:
        if isinstance(analysis, AnalysisReport):
            return analysis
        report = coerce_analysis_report(analysis)
        return AnalysisReport(
            title=report.title,
            summary=report.summary,
            sections=report.sections,
            metadata=report.metadata,
            language=language,
        )

    def _load_module(self, path: Path, module_name: str) -> ModuleType:
        spec = spec_from_file_location(module_name, path)
        if spec is None or spec.loader is None:
            raise ImportError(f"Unable to load module at {path}")
        module = module_from_spec(spec)
        spec.loader.exec_module(module)
        return module

    def _resolve_language(self, analysis: AnalysisResult | AnalysisReport, config: ReportOutputConfig) -> str:
        if isinstance(analysis, AnalysisReport):
            return normalize_report_language(config.language or analysis.language)
        return normalize_report_language(config.language)

    def _resolve_text_map(self, mapping: dict[str, dict[str, object]], language: str) -> dict[str, object]:
        return dict(mapping.get(language) or mapping["zh-CN"])


def _major_heading_prefix(index: int, *, language: str) -> str:
    if language == "en-US":
        return f"{_to_roman(index)}. "
    return f"{_to_chinese_ordinal(index)}、"


def _to_chinese_ordinal(value: int) -> str:
    numerals = {
        0: "零",
        1: "一",
        2: "二",
        3: "三",
        4: "四",
        5: "五",
        6: "六",
        7: "七",
        8: "八",
        9: "九",
        10: "十",
    }
    if value <= 10:
        return numerals[value]
    if value < 20:
        return f"十{numerals[value - 10]}"
    tens, ones = divmod(value, 10)
    tens_text = numerals[tens]
    if ones == 0:
        return f"{tens_text}十"
    return f"{tens_text}十{numerals[ones]}"


def _to_roman(value: int) -> str:
    pairs = (
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    )
    remaining = value
    parts: list[str] = []
    for number, numeral in pairs:
        while remaining >= number:
            parts.append(numeral)
            remaining -= number
    return "".join(parts)
