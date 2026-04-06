from __future__ import annotations

from dataclasses import dataclass
from types import ModuleType

from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


@dataclass(frozen=True)
class FontSpec:
    latin: str
    east_asia: str
    size_pt: float
    bold: bool = False


@dataclass(frozen=True)
class ParagraphSpec:
    style_name: str
    font: FontSpec
    alignment: WD_ALIGN_PARAGRAPH | None = None
    space_before_pt: float = 0
    space_after_pt: float = 0
    line_spacing: float = 1.35


def apply_document_theme(document: Document, *, language: str) -> None:
    _apply_page_layout(document)
    theme = _theme_for(language)

    _configure_paragraph_style(document, theme["cover_title"])
    _configure_paragraph_style(document, theme["cover_subtitle"])
    _configure_paragraph_style(document, theme["heading_1"])
    _configure_paragraph_style(document, theme["heading_2"])
    _configure_paragraph_style(document, theme["body"])
    _configure_paragraph_style(document, theme["table_title"])
    _configure_character_style(document, "DBA Table Header", theme["table_header"].font)
    _configure_character_style(document, "DBA Table Body", theme["table_body"].font)

    normal = document.styles["Normal"]
    _apply_font(normal.font, theme["body"].font)


def add_cover_title(document: Document, text: str) -> None:
    document.add_paragraph(text, style="DBA Cover Title")


def add_cover_subtitle(document: Document, text: str) -> None:
    if text:
        document.add_paragraph(text, style="DBA Cover Subtitle")


def add_cover_metadata(document: Document, label: str, value: str, *, language: str) -> None:
    paragraph = document.add_paragraph(style="DBA Cover Subtitle")
    run = paragraph.add_run(f"{label}: {value}")
    _apply_run_font(run, _theme_for(language)["cover_subtitle"].font)


def add_page_break(document: Document) -> None:
    paragraph = document.add_paragraph(style="DBA Body Text")
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_heading(document: Document, text: str, *, level: int) -> None:
    style_name = "DBA Heading 1" if level <= 1 else "DBA Heading 2"
    document.add_paragraph(text, style=style_name)


def add_body_paragraph(document: Document, text: str) -> None:
    document.add_paragraph(text, style="DBA Body Text")


def add_table_title(document: Document, text: str) -> None:
    document.add_paragraph(text, style="DBA Table Title")


def style_table(table, *, language: str, table_style_module: ModuleType) -> None:
    theme = _theme_for(language)
    table.style = table_style_module.DEFAULT_DOCX_TABLE_STYLE
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_fixed_layout(table)
    _apply_column_widths(table, table_style_module)

    for row_index, row in enumerate(table.rows):
        is_header = row_index == 0
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    _apply_run_font(run, theme["table_header"].font if is_header else theme["table_body"].font)
                if is_header:
                    _apply_cell_shading(cell, table_style_module.DOCX_TABLE_HEADER_FILL)


def _apply_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)


def _theme_for(language: str) -> dict[str, ParagraphSpec]:
    if language == "en-US":
        serif = "Times New Roman"
        sans = "Calibri"
        return {
            "cover_title": ParagraphSpec("DBA Cover Title", FontSpec(latin=sans, east_asia="Cambria", size_pt=18, bold=True), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before_pt=36, space_after_pt=12),
            "cover_subtitle": ParagraphSpec("DBA Cover Subtitle", FontSpec(latin=serif, east_asia="Cambria", size_pt=12), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=6),
            "heading_1": ParagraphSpec("DBA Heading 1", FontSpec(latin=sans, east_asia="Cambria", size_pt=14, bold=True), space_before_pt=14, space_after_pt=8),
            "heading_2": ParagraphSpec("DBA Heading 2", FontSpec(latin=sans, east_asia="Cambria", size_pt=12, bold=True), space_before_pt=10, space_after_pt=6),
            "body": ParagraphSpec("DBA Body Text", FontSpec(latin=serif, east_asia="Cambria", size_pt=12), space_after_pt=6),
            "table_title": ParagraphSpec("DBA Table Title", FontSpec(latin=sans, east_asia="Cambria", size_pt=11, bold=True), space_before_pt=6, space_after_pt=4),
            "table_header": ParagraphSpec("DBA Table Header", FontSpec(latin=serif, east_asia="Cambria", size_pt=10, bold=True)),
            "table_body": ParagraphSpec("DBA Table Body", FontSpec(latin=serif, east_asia="Cambria", size_pt=10)),
        }
    return {
        "cover_title": ParagraphSpec("DBA Cover Title", FontSpec(latin="SimHei", east_asia="黑体", size_pt=18, bold=True), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before_pt=36, space_after_pt=12),
        "cover_subtitle": ParagraphSpec("DBA Cover Subtitle", FontSpec(latin="SimSun", east_asia="宋体", size_pt=12), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=6),
        "heading_1": ParagraphSpec("DBA Heading 1", FontSpec(latin="SimHei", east_asia="黑体", size_pt=14, bold=True), space_before_pt=14, space_after_pt=8),
        "heading_2": ParagraphSpec("DBA Heading 2", FontSpec(latin="SimHei", east_asia="黑体", size_pt=12, bold=True), space_before_pt=10, space_after_pt=6),
        "body": ParagraphSpec("DBA Body Text", FontSpec(latin="SimSun", east_asia="宋体", size_pt=12), space_after_pt=6),
        "table_title": ParagraphSpec("DBA Table Title", FontSpec(latin="SimSun", east_asia="宋体", size_pt=10.5, bold=True), space_before_pt=6, space_after_pt=4),
        "table_header": ParagraphSpec("DBA Table Header", FontSpec(latin="SimSun", east_asia="宋体", size_pt=10.5, bold=True)),
        "table_body": ParagraphSpec("DBA Table Body", FontSpec(latin="SimSun", east_asia="宋体", size_pt=10.5)),
    }


def _configure_paragraph_style(document: Document, spec: ParagraphSpec) -> None:
    style = document.styles[spec.style_name] if spec.style_name in document.styles else document.styles.add_style(spec.style_name, WD_STYLE_TYPE.PARAGRAPH)
    _apply_font(style.font, spec.font)
    style.paragraph_format.space_before = Pt(spec.space_before_pt)
    style.paragraph_format.space_after = Pt(spec.space_after_pt)
    style.paragraph_format.line_spacing = spec.line_spacing
    if spec.alignment is not None:
        style.paragraph_format.alignment = spec.alignment


def _configure_character_style(document: Document, style_name: str, font: FontSpec) -> None:
    style = document.styles[style_name] if style_name in document.styles else document.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
    _apply_font(style.font, font)


def _apply_font(font, spec: FontSpec) -> None:
    font.name = spec.latin
    font.size = Pt(spec.size_pt)
    font.bold = spec.bold
    r_fonts = font.element.rPr.rFonts
    r_fonts.set(qn("w:ascii"), spec.latin)
    r_fonts.set(qn("w:hAnsi"), spec.latin)
    r_fonts.set(qn("w:eastAsia"), spec.east_asia)
    r_fonts.set(qn("w:cs"), spec.latin)


def _apply_run_font(run, spec: FontSpec) -> None:
    run.bold = spec.bold
    font = run.font
    font.name = spec.latin
    font.size = Pt(spec.size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:ascii"), spec.latin)
    r_fonts.set(qn("w:hAnsi"), spec.latin)
    r_fonts.set(qn("w:eastAsia"), spec.east_asia)
    r_fonts.set(qn("w:cs"), spec.latin)


def _set_fixed_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def _apply_column_widths(table, table_style_module: ModuleType) -> None:
    if not table.rows:
        return
    column_count = len(table.columns)
    lengths = [0] * column_count
    for row in table.rows[: min(len(table.rows), 16)]:
        for index, cell in enumerate(row.cells):
            lengths[index] = max(lengths[index], len(cell.text.strip()))
    weighted = [max(8, min(length, 28)) for length in lengths]
    total_weight = sum(weighted) or column_count
    total_width = table_style_module.DOCX_TABLE_TOTAL_WIDTH_INCHES

    widths = []
    for weight in weighted:
        width = total_width * (weight / total_weight)
        width = max(table_style_module.DOCX_TABLE_MIN_COL_WIDTH_INCHES, width)
        width = min(table_style_module.DOCX_TABLE_MAX_COL_WIDTH_INCHES, width)
        widths.append(Inches(width))

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = widths[index]


def _apply_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
