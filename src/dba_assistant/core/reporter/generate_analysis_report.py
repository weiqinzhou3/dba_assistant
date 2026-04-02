from __future__ import annotations

from typing import Any

from docx import Document

from dba_assistant.core.analyzer.types import AnalysisResult, ReportSection, TableModel
from dba_assistant.core.reporter.report_model import AnalysisReport, ReportSectionModel, TableBlock, TextBlock
from dba_assistant.core.reporter.types import ReportArtifact, ReportFormat, ReportOutputConfig


def generate_analysis_report(
    analysis: AnalysisReport | AnalysisResult,
    config: ReportOutputConfig,
) -> ReportArtifact:
    report = _coerce_report_model(analysis)

    if config.format is ReportFormat.SUMMARY:
        content = render_summary_text(report)
        if config.output_path is not None:
            config.output_path.parent.mkdir(parents=True, exist_ok=True)
            config.output_path.write_text(content, encoding="utf-8")
        return ReportArtifact(format=ReportFormat.SUMMARY, output_path=config.output_path, content=content)

    if config.format is ReportFormat.DOCX:
        if config.output_path is None:
            raise ValueError("Docx report generation requires an output_path.")
        document = Document()
        render_docx_report(document, report)
        config.output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(config.output_path)
        return ReportArtifact(format=ReportFormat.DOCX, output_path=config.output_path, content=None)

    raise NotImplementedError(f"Unsupported report format: {config.format}")


def render_summary_text(report: AnalysisReport) -> str:
    lines = [report.title]

    if report.summary:
        lines.extend(["", report.summary])

    if report.metadata:
        lines.extend(["", "Metadata"])
        for key, value in sorted(report.metadata.items()):
            lines.append(f"- {key}: {value}")

    for section in report.sections:
        lines.extend(["", section.title])
        for block in section.blocks:
            if isinstance(block, TextBlock):
                lines.append(block.text)
                continue
            lines.append(block.title)
            if block.columns:
                lines.append(", ".join(block.columns))
            for row in block.rows:
                lines.append(", ".join(row))

    return "\n".join(lines)


def render_docx_report(document: Document, report: AnalysisReport, *, table_style: str = "Table Grid") -> None:
    document.add_heading(report.title, level=0)

    if report.summary:
        document.add_paragraph(report.summary)

    if report.metadata:
        document.add_heading("Metadata", level=1)
        for key, value in sorted(report.metadata.items()):
            document.add_paragraph(f"{key}: {value}")

    render_docx_sections(document, report.sections, table_style=table_style)


def render_docx_sections(
    document: Document,
    sections: list[ReportSectionModel],
    *,
    table_style: str = "Table Grid",
) -> None:
    for section in sections:
        document.add_heading(section.title, level=1)
        for block in section.blocks:
            if isinstance(block, TextBlock):
                document.add_paragraph(block.text)
                continue
            document.add_paragraph(block.title)
            docx_table = document.add_table(rows=1, cols=len(block.columns))
            docx_table.style = table_style
            for index, column in enumerate(block.columns):
                docx_table.rows[0].cells[index].text = column
            for row in block.rows:
                cells = docx_table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = value


def _coerce_report_model(analysis: AnalysisReport | AnalysisResult) -> AnalysisReport:
    if isinstance(analysis, AnalysisReport):
        return analysis

    sections: list[ReportSectionModel] = []

    if analysis.risk_summary:
        sections.append(
            ReportSectionModel(
                id="risk_summary",
                title="Risk Summary",
                blocks=[
                    TextBlock(text=f"- {key}: {value}")
                    for key, value in sorted(analysis.risk_summary.items())
                ],
            )
        )

    for section in analysis.sections:
        sections.append(_coerce_section(section))

    return AnalysisReport(
        title=analysis.title,
        summary=analysis.summary,
        sections=sections,
        metadata={key: str(value) for key, value in analysis.metadata.items()},
    )


def _coerce_section(section: ReportSection) -> ReportSectionModel:
    blocks: list[TextBlock | TableBlock] = []

    if section.summary:
        blocks.append(TextBlock(text=section.summary))

    for paragraph in section.paragraphs:
        blocks.append(TextBlock(text=paragraph))

    for table in section.tables:
        blocks.append(_coerce_table(table))

    return ReportSectionModel(
        id=_section_id(section.title),
        title=section.title,
        blocks=blocks,
    )


def _coerce_table(table: TableModel) -> TableBlock:
    return TableBlock(
        title=table.title,
        columns=[str(column) for column in table.columns],
        rows=[[str(cell) for cell in row] for row in table.rows],
    )


def _section_id(title: str) -> str:
    return title.lower().replace(" ", "_")
