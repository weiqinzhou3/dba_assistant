import json
from pathlib import Path

import pytest
from docx import Document

from dba_assistant.core.reporter.docx_reporter import DocxReporter
from dba_assistant.core.reporter.report_model import AnalysisReport
from dba_assistant.core.reporter.report_model import render_summary_text
from dba_assistant.core.reporter.types import OutputMode, ReportFormat, ReportOutputConfig
from dba_assistant.adaptors.mysql_adaptor import MySQLConnectionConfig
from dba_assistant.parsers import rdb_parser_strategy as parser_strategy_module
from dba_assistant.parsers.rdb_parser_strategy import StreamedRowsResult
from dba_assistant.capabilities.redis_rdb_analysis import service as service_module
from dba_assistant.capabilities.redis_rdb_analysis.analyzers.overall import analyze_overall
from dba_assistant.capabilities.redis_rdb_analysis.analyzers.rcs_custom import analyze_rcs_custom
from dba_assistant.capabilities.redis_rdb_analysis.service import analyze_rdb
from dba_assistant.capabilities.redis_rdb_analysis.types import (
    AnalysisStatus,
    ConfirmationRequest,
    InputSourceKind,
    KeyRecord,
    NormalizedRdbDataset,
    RdbAnalysisRequest,
    SampleInput,
)
from dba_assistant.tools.mysql_tools import (
    load_preparsed_dataset_from_mysql,
    stage_rdb_rows_to_mysql,
)

HDT_BINARY = Path(".tools/bin/rdb")


class InMemoryTextMySQLAdaptor:
    def __init__(self) -> None:
        self._rows_by_table: dict[str, list[dict[str, object]]] = {}

    def execute_write(self, _config, sql: str, params=None) -> int:
        if sql.startswith("CREATE TABLE IF NOT EXISTS "):
            table_name = sql.removeprefix("CREATE TABLE IF NOT EXISTS ").split(" ", 1)[0].strip("`")
            self._rows_by_table.setdefault(table_name, [])
            return 0

        if sql.startswith("INSERT INTO "):
            table_name = sql.removeprefix("INSERT INTO ").split(" ", 1)[0].strip("`")
            column_names = [
                column.strip().strip("`")
                for column in sql.split("(", 1)[1].split(")", 1)[0].split(",")
            ]
            stored_rows = self._rows_by_table.setdefault(table_name, [])
            for row_values in params or []:
                stored_rows.append(
                    {
                        column: None if value is None else str(value)
                        for column, value in zip(column_names, row_values, strict=True)
                    }
                )
            return len(params or [])

        raise AssertionError(f"Unexpected write SQL: {sql}")

    def read_query(self, _config, sql: str) -> list[dict[str, object]]:
        if not sql.startswith("SELECT * FROM "):
            raise AssertionError(f"Unexpected read SQL: {sql}")

        tail = sql.removeprefix("SELECT * FROM ")
        table_name, _, raw_limit = tail.partition(" LIMIT ")
        rows = self._rows_by_table.get(table_name.strip("`"), [])
        return [dict(row) for row in rows[: int(raw_limit)]]


def _build_mysql_side_analysis(rows: list[dict[str, object]]):
    def analyze_staged(staging, *, profile, sample_rows):
        dataset = NormalizedRdbDataset(
            samples=[
                SampleInput(
                    source=row[2],
                    kind=InputSourceKind.LOCAL_RDB,
                    label=row[0],
                )
                for row in sample_rows
            ],
            records=[
                KeyRecord(
                    sample_id="sample-1",
                    key_name=str(row["key_name"]),
                    key_type=str(row["key_type"]),
                    size_bytes=int(row["size_bytes"]),
                    has_expiration=bool(row["has_expiration"]),
                    ttl_seconds=None if row["ttl_seconds"] is None else int(row["ttl_seconds"]),
                    prefix_segments=tuple(part for part in str(row["key_name"]).split(":")[:-1] if part),
                )
                for row in rows
            ],
        )
        result = analyze_overall(dataset, profile=profile)
        if profile.name.lower() == "rcs":
            result.update(analyze_rcs_custom(dataset))
        return result

    return analyze_staged


def _streamed_rows(rows: list[dict[str, object]]) -> StreamedRowsResult:
    return StreamedRowsResult(rows=iter(rows), strategy_name="test-stream")


def test_analyze_rdb_returns_confirmation_request_for_remote_redis_without_confirmation() -> None:
    request = RdbAnalysisRequest(
        prompt="analyze latest rdb",
        inputs=[SampleInput(source="10.0.0.8:6379", kind=InputSourceKind.REMOTE_REDIS)],
    )

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
    )

    assert isinstance(result, ConfirmationRequest)
    assert result.status is AnalysisStatus.CONFIRMATION_REQUIRED
    assert result.required_action == "fetch_existing"
    assert "/data/redis/dump.rdb" in result.message


def test_analyze_rdb_returns_analysis_report_for_local_inputs(monkeypatch) -> None:
    rows = json.loads(Path("tests/fixtures/rdb/direct/sample_key_records.json").read_text(encoding="utf-8"))
    request = RdbAnalysisRequest(
        prompt="analyze this rdb",
        inputs=[SampleInput(source=Path("/tmp/dump.rdb"), kind=InputSourceKind.LOCAL_RDB)],
    )
    monkeypatch.setattr("dba_assistant.capabilities.redis_rdb_analysis.service._parse_rdb_rows", lambda _path: rows)

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
    )

    assert isinstance(result, AnalysisReport)
    assert result.title == "Redis RDB 分析报告"
    assert result.metadata["profile"] == "generic"
    assert result.metadata["route"] == "direct_rdb_analysis"
    assert result.metadata["path"] == "3c"
    assert any(section.id == "top_big_keys" for section in result.sections)


def test_analyze_rdb_applies_requested_rcs_profile_top_n_and_focus_prefixes(monkeypatch) -> None:
    rows = [
        {"key_name": "order:1", "key_type": "string", "size_bytes": 900, "has_expiration": True, "ttl_seconds": 60},
        {"key_name": "order:2", "key_type": "hash", "size_bytes": 800, "has_expiration": False, "ttl_seconds": None},
        {"key_name": "order:3", "key_type": "string", "size_bytes": 700, "has_expiration": False, "ttl_seconds": None},
        {"key_name": "mq:1", "key_type": "stream", "size_bytes": 600, "has_expiration": False, "ttl_seconds": None},
        {"key_name": "mq:2", "key_type": "stream", "size_bytes": 500, "has_expiration": True, "ttl_seconds": 30},
        {"key_name": "loan:1", "key_type": "hash", "size_bytes": 400, "has_expiration": False, "ttl_seconds": None},
    ]
    request = RdbAnalysisRequest(
        prompt="使用 rcs profile，只输出 top 2，只看 order:* 和 mq:*",
        inputs=[SampleInput(source=Path("/tmp/dump.rdb"), kind=InputSourceKind.LOCAL_RDB)],
        profile_name="rcs",
        profile_overrides={
            "top_n": {
                "prefix_top": 2,
                "top_big_keys": 2,
                "string_big_keys": 2,
                "hash_big_keys": 2,
                "list_big_keys": 2,
                "set_big_keys": 2,
                "zset_big_keys": 2,
                "stream_big_keys": 2,
                "other_big_keys": 2,
                "focused_prefix_top_keys": 2,
            },
            "focus_prefixes": ("order:*", "mq:*"),
        },
    )
    monkeypatch.setattr("dba_assistant.capabilities.redis_rdb_analysis.service._parse_rdb_rows", lambda _path: rows)

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
    )

    assert isinstance(result, AnalysisReport)
    assert result.metadata["profile"] == "rcs"
    assert all(section.id != "sample_overview" for section in result.sections)
    assert any(section.id == "focused_prefix_analysis" for section in result.sections)
    assert any(section.title == "重点前缀详情分析" for section in result.sections)
    top_big_key_section = next(section for section in result.sections if section.id == "top_big_keys")
    assert len(top_big_key_section.blocks[1].rows) == 2
    focused_order_section = next(section for section in result.sections if section.title == "前缀 order:* 详情")
    assert "loan:*" not in focused_order_section.blocks[0].text
    assert len(focused_order_section.blocks[1].rows) == 2


def test_analyze_rdb_focus_only_mode_uses_only_focused_prefix_sections(monkeypatch) -> None:
    rows = [
        {"key_name": "tag:1", "key_type": "string", "size_bytes": 900, "has_expiration": True, "ttl_seconds": 60},
        {"key_name": "tag:2", "key_type": "hash", "size_bytes": 800, "has_expiration": False, "ttl_seconds": None},
        {"key_name": "loan:1", "key_type": "hash", "size_bytes": 400, "has_expiration": False, "ttl_seconds": None},
    ]
    request = RdbAnalysisRequest(
        prompt="只需要输出前缀为tag的key，其他都不需要",
        inputs=[SampleInput(source=Path("/tmp/dump.rdb"), kind=InputSourceKind.LOCAL_RDB)],
        profile_name="rcs",
        profile_overrides={
            "focus_prefixes": ("tag:*",),
            "focus_only": True,
            "top_n": {
                "prefix_top": 10,
                "focused_prefix_top_keys": 10,
                "top_big_keys": 10,
                "string_big_keys": 10,
                "hash_big_keys": 10,
                "list_big_keys": 10,
                "set_big_keys": 10,
                "zset_big_keys": 10,
                "stream_big_keys": 10,
                "other_big_keys": 10,
            },
        },
    )
    monkeypatch.setattr("dba_assistant.capabilities.redis_rdb_analysis.service._parse_rdb_rows", lambda _path: rows)

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
    )

    assert isinstance(result, AnalysisReport)
    assert result.metadata["scope"] == "focused_prefix_only"
    assert [section.title for section in result.sections] == ["重点前缀详情分析", "前缀 tag:* 详情"]


def test_analyze_rdb_focus_only_mode_keeps_top_10_in_summary_and_docx(tmp_path: Path, monkeypatch) -> None:
    rows = [
        {
            "key_name": f"signin:{index}",
            "key_type": "string" if index % 2 else "hash",
            "size_bytes": 1000 - index,
            "has_expiration": bool(index % 2),
            "ttl_seconds": 60 if index % 2 else None,
        }
        for index in range(1, 13)
    ]
    request = RdbAnalysisRequest(
        prompt="前缀是signin前缀的key，只需要top 10 其他分析结果都不需要",
        inputs=[SampleInput(source=Path("/tmp/dump.rdb"), kind=InputSourceKind.LOCAL_RDB)],
        profile_name="rcs",
        profile_overrides={
            "focus_prefixes": ("signin:*",),
            "focus_only": True,
            "top_n": {
                "prefix_top": 10,
                "focused_prefix_top_keys": 10,
                "top_big_keys": 10,
                "string_big_keys": 10,
                "hash_big_keys": 10,
                "list_big_keys": 10,
                "set_big_keys": 10,
                "zset_big_keys": 10,
                "stream_big_keys": 10,
                "other_big_keys": 10,
            },
        },
    )
    monkeypatch.setattr("dba_assistant.capabilities.redis_rdb_analysis.service._parse_rdb_rows", lambda _path: rows)

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
    )

    assert isinstance(result, AnalysisReport)
    assert result.metadata["scope"] == "focused_prefix_only"
    detail = next(section for section in result.sections if section.title == "前缀 signin:* 详情")
    assert detail.blocks[1].title == "前缀 signin:* Top Keys（Top 10）"
    assert len(detail.blocks[1].rows) == 10

    summary_text = render_summary_text(result, language="zh-CN")
    assert "前缀 signin:* Top Keys（Top 10）" in summary_text

    output_path = tmp_path / "focus-signin-top10.docx"
    DocxReporter().render(
        result,
        ReportOutputConfig(
            output_path=output_path,
            mode=OutputMode.REPORT,
            format=ReportFormat.DOCX,
            template_name="rdb-analysis",
            language="zh-CN",
        ),
    )
    document_text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)
    assert "前缀 signin:* Top Keys（Top 10）" in document_text


def test_analyze_rdb_full_report_keeps_standard_sections_and_adds_requested_prefix_details(monkeypatch) -> None:
    rows = [
        {"key_name": "tag:1", "key_type": "string", "size_bytes": 900, "has_expiration": True, "ttl_seconds": 60},
        {"key_name": "tag:2", "key_type": "hash", "size_bytes": 800, "has_expiration": False, "ttl_seconds": None},
        {"key_name": "store:1", "key_type": "string", "size_bytes": 700, "has_expiration": False, "ttl_seconds": None},
        {"key_name": "store:2", "key_type": "hash", "size_bytes": 600, "has_expiration": True, "ttl_seconds": 30},
        {"key_name": "loan:1", "key_type": "hash", "size_bytes": 500, "has_expiration": False, "ttl_seconds": None},
    ]
    request = RdbAnalysisRequest(
        prompt="使用 rcs profile，重点分析 tag 和 store，top 2",
        inputs=[SampleInput(source=Path("/tmp/dump.rdb"), kind=InputSourceKind.LOCAL_RDB)],
        profile_name="rcs",
        profile_overrides={
            "focus_prefixes": ("tag:*", "store:*"),
            "top_n": {
                "prefix_top": 2,
                "focused_prefix_top_keys": 2,
                "top_big_keys": 2,
                "string_big_keys": 2,
                "hash_big_keys": 2,
                "list_big_keys": 2,
                "set_big_keys": 2,
                "zset_big_keys": 2,
                "stream_big_keys": 2,
                "other_big_keys": 2,
            },
        },
    )
    monkeypatch.setattr("dba_assistant.capabilities.redis_rdb_analysis.service._parse_rdb_rows", lambda _path: rows)

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
    )

    assert isinstance(result, AnalysisReport)
    assert result.metadata["profile"] == "rcs"
    assert result.metadata["scope"] == "full_report"
    assert any(section.id == "top_big_keys" for section in result.sections)
    assert [section.title for section in result.sections if section.id.startswith("focused_prefix")] == [
        "重点前缀详情分析",
        "前缀 tag:* 详情",
        "前缀 store:* 详情",
    ]
    tag_section = next(section for section in result.sections if section.title == "前缀 tag:* 详情")
    store_section = next(section for section in result.sections if section.title == "前缀 store:* 详情")
    assert len(tag_section.blocks[1].rows) == 2
    assert len(store_section.blocks[1].rows) == 2


def test_analyze_rdb_full_report_keeps_standard_sections_and_applies_top_10_to_uv_details(tmp_path: Path, monkeypatch) -> None:
    rows = [
        {
            "key_name": f"uv:{index}",
            "key_type": "string",
            "size_bytes": 1000 - index,
            "has_expiration": False,
            "ttl_seconds": None,
        }
        for index in range(1, 13)
    ] + [
        {"key_name": "other:1", "key_type": "hash", "size_bytes": 100, "has_expiration": True, "ttl_seconds": 60}
    ]
    request = RdbAnalysisRequest(
        prompt="使用 rcs profile，重点分析前缀为 uv 的 key，top 10",
        inputs=[SampleInput(source=Path("/tmp/dump.rdb"), kind=InputSourceKind.LOCAL_RDB)],
        profile_name="rcs",
        profile_overrides={
            "focus_prefixes": ("uv:*",),
            "top_n": {
                "prefix_top": 10,
                "focused_prefix_top_keys": 10,
                "top_big_keys": 10,
                "string_big_keys": 10,
                "hash_big_keys": 10,
                "list_big_keys": 10,
                "set_big_keys": 10,
                "zset_big_keys": 10,
                "stream_big_keys": 10,
                "other_big_keys": 10,
            },
        },
    )
    monkeypatch.setattr("dba_assistant.capabilities.redis_rdb_analysis.service._parse_rdb_rows", lambda _path: rows)

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
    )

    assert isinstance(result, AnalysisReport)
    assert result.metadata["scope"] == "full_report"
    assert any(section.id == "top_big_keys" for section in result.sections)
    detail = next(section for section in result.sections if section.title == "前缀 uv:* 详情")
    assert detail.blocks[1].title == "前缀 uv:* Top Keys（Top 10）"
    assert len(detail.blocks[1].rows) == 10

    summary_text = render_summary_text(result, language="zh-CN")
    assert "总体大 Key 排名（Top 10）" in summary_text
    assert "前缀 uv:* Top Keys（Top 10）" in summary_text

    output_path = tmp_path / "full-uv-top10.docx"
    DocxReporter().render(
        result,
        ReportOutputConfig(
            output_path=output_path,
            mode=OutputMode.REPORT,
            format=ReportFormat.DOCX,
            template_name="rdb-analysis",
            language="zh-CN",
        ),
    )
    document_text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)
    assert "总体大 Key 排名（Top 10）" in document_text
    assert "前缀 uv:* Top Keys（Top 10）" in document_text


def test_analyze_rdb_mysql_backed_prefix_detail_uses_canonical_dataset_without_value_size_dependency(
    monkeypatch,
) -> None:
    rows = [
        {"key_name": "session:data:1", "key_type": "string", "size_bytes": 123, "has_expiration": False, "ttl_seconds": None},
        {"key_name": "session:data:2", "key_type": "hash", "size_bytes": 88, "has_expiration": True, "ttl_seconds": 60},
        {"key_name": "other:1", "key_type": "string", "size_bytes": 77, "has_expiration": False, "ttl_seconds": None},
    ]
    request = RdbAnalysisRequest(
        prompt="只输出 session:data 的 key 详情",
        inputs=[SampleInput(source=Path("/tmp/dump.rdb"), kind=InputSourceKind.LOCAL_RDB)],
        profile_name="rcs",
        path_mode="database_backed_analysis",
        profile_overrides={
            "focus_prefixes": ("session:data:*",),
            "focus_only": True,
        },
    )
    calls: dict[str, object] = {}
    monkeypatch.setattr(
        "dba_assistant.capabilities.redis_rdb_analysis.service._stream_rdb_rows",
        lambda _path: _streamed_rows(rows),
    )

    def fake_stage_rdb_rows_to_mysql(
        table_name: str,
        parsed_rows: list[dict[str, object]],
        *,
        source_file: str = "manual",
        run_id: str = "manual",
    ) -> dict[str, object]:
        calls["stage"] = (table_name, parsed_rows)
        return {"table": table_name, "staged": len(parsed_rows), "run_id": run_id}

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
        stage_rdb_rows_to_mysql=fake_stage_rdb_rows_to_mysql,
        analyze_staged_rdb_rows=_build_mysql_side_analysis(rows),
    )

    assert isinstance(result, AnalysisReport)
    assert result.metadata["route"] == "database_backed_analysis"
    assert result.metadata["scope"] == "focused_prefix_only"
    assert [section.title for section in result.sections] == ["重点前缀详情分析", "前缀 session:data:* 详情"]


@pytest.mark.parametrize(
    ("focus_prefix", "expected_title", "expected_rows"),
    (
        ("tag:*", "前缀 tag:* 详情", 1),
        ("signin:*", "前缀 signin:* 详情", 1),
        ("store:*", "前缀 store:* 详情", 1),
        ("uv:*", "前缀 uv:* 详情", 1),
        ("session:data:*", "前缀 session:data:* 详情", 1),
    ),
)
def test_analyze_rdb_mysql_backed_prefix_details_support_arbitrary_prefixes_without_mysql_column_assumptions(
    monkeypatch,
    focus_prefix: str,
    expected_title: str,
    expected_rows: int,
) -> None:
    rows = [
        {"key_name": "tag:1", "key_type": "string", "size_bytes": 120, "has_expiration": False, "ttl_seconds": None},
        {"key_name": "signin:1", "key_type": "hash", "size_bytes": 110, "has_expiration": True, "ttl_seconds": 60},
        {"key_name": "store:1", "key_type": "string", "size_bytes": 100, "has_expiration": False, "ttl_seconds": None},
        {"key_name": "uv:1", "key_type": "string", "size_bytes": 95, "has_expiration": False, "ttl_seconds": None},
        {"key_name": "session:data:1", "key_type": "string", "size_bytes": 90, "has_expiration": False, "ttl_seconds": None},
    ]
    request = RdbAnalysisRequest(
        prompt=f"只输出 {focus_prefix} 的 key 详情",
        inputs=[SampleInput(source=Path("/tmp/dump.rdb"), kind=InputSourceKind.LOCAL_RDB)],
        profile_name="rcs",
        path_mode="database_backed_analysis",
        profile_overrides={
            "focus_prefixes": (focus_prefix,),
            "focus_only": True,
            "top_n": {"focused_prefix_top_keys": 10},
        },
    )
    monkeypatch.setattr(
        "dba_assistant.capabilities.redis_rdb_analysis.service._stream_rdb_rows",
        lambda _path: _streamed_rows(rows),
    )

    def fake_stage_rdb_rows_to_mysql(
        table_name: str,
        parsed_rows: list[dict[str, object]],
        *,
        source_file: str = "manual",
        run_id: str = "manual",
    ) -> dict[str, object]:
        return {"table": table_name, "staged": len(parsed_rows), "run_id": run_id}

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
        stage_rdb_rows_to_mysql=fake_stage_rdb_rows_to_mysql,
        analyze_staged_rdb_rows=_build_mysql_side_analysis(rows),
    )

    assert isinstance(result, AnalysisReport)
    assert result.metadata["route"] == "database_backed_analysis"
    assert result.metadata["scope"] == "focused_prefix_only"
    detail = next(section for section in result.sections if section.title == expected_title)
    assert len(detail.blocks[1].rows) == expected_rows


def test_analyze_rdb_mysql_backed_prefix_detail_keeps_zero_match_section(monkeypatch) -> None:
    rows = [
        {"key_name": "tag:1", "key_type": "string", "size_bytes": 123, "has_expiration": False, "ttl_seconds": None},
    ]
    request = RdbAnalysisRequest(
        prompt="只输出 signin 的 key 详情",
        inputs=[SampleInput(source="mysql:preparsed_keys", kind=InputSourceKind.PREPARSED_MYSQL)],
        profile_name="rcs",
        mysql_table="preparsed_keys",
        profile_overrides={
            "focus_prefixes": ("signin:*",),
            "focus_only": True,
        },
    )

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
        load_preparsed_dataset_from_mysql=lambda _table_name: {"source": "mysql:preparsed_keys", "rows": rows},
    )

    assert isinstance(result, AnalysisReport)
    assert any(section.title == "前缀 signin:* 详情" for section in result.sections)
    detail = next(section for section in result.sections if section.title == "前缀 signin:* 详情")
    assert detail.blocks[0].text == "未匹配到符合条件的键。"


def test_analyze_rdb_local_inputs_do_not_call_remote_discovery(monkeypatch) -> None:
    rows = json.loads(Path("tests/fixtures/rdb/direct/sample_key_records.json").read_text(encoding="utf-8"))
    request = RdbAnalysisRequest(
        prompt="analyze this local rdb",
        inputs=[SampleInput(source=Path("/tmp/dump.rdb"), kind=InputSourceKind.LOCAL_RDB)],
    )
    monkeypatch.setattr("dba_assistant.capabilities.redis_rdb_analysis.service._parse_rdb_rows", lambda _path: rows)

    def fail_remote_discovery(*_args, **_kwargs):
        raise AssertionError("remote_discovery should not be called for local inputs")

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=fail_remote_discovery,
    )

    assert isinstance(result, AnalysisReport)


def test_analyze_rdb_returns_analysis_report_for_precomputed_inputs() -> None:
    request = RdbAnalysisRequest(
        prompt="summarize this exported analysis",
        inputs=[
            SampleInput(
                source=Path("tests/fixtures/rdb/precomputed/sample_precomputed_rows.json"),
                kind=InputSourceKind.PRECOMPUTED,
            )
        ],
    )

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
    )

    assert isinstance(result, AnalysisReport)
    assert result.title == "Redis RDB 分析报告"
    assert result.metadata["profile"] == "generic"
    assert result.metadata["route"] == "preparsed_dataset_analysis"
    assert result.metadata["path"] == "3b"
    assert any(section.id == "sample_overview" for section in result.sections)


def test_analyze_rdb_returns_analysis_report_for_preparsed_mysql_inputs() -> None:
    rows = json.loads(Path("tests/fixtures/rdb/direct/sample_key_records.json").read_text(encoding="utf-8"))
    request = RdbAnalysisRequest(
        prompt="summarize this mysql dataset",
        inputs=[SampleInput(source="mysql:preparsed_keys", kind=InputSourceKind.PREPARSED_MYSQL)],
        mysql_table="preparsed_keys",
    )
    calls: dict[str, object] = {}

    def fake_load_preparsed_dataset_from_mysql(table_name: str) -> dict[str, object]:
        calls["table_name"] = table_name
        return {"source": f"mysql:{table_name}", "rows": rows}

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
        load_preparsed_dataset_from_mysql=fake_load_preparsed_dataset_from_mysql,
    )

    assert isinstance(result, AnalysisReport)
    assert calls["table_name"] == "preparsed_keys"
    assert result.metadata["route"] == "preparsed_dataset_analysis"
    assert result.metadata["path"] == "3b"
    assert any(section.id == "sample_overview" for section in result.sections)


def test_analyze_rdb_supports_explicit_english_report_language(monkeypatch) -> None:
    rows = json.loads(Path("tests/fixtures/rdb/direct/sample_key_records.json").read_text(encoding="utf-8"))
    request = RdbAnalysisRequest(
        prompt="Analyze this rdb in English",
        inputs=[SampleInput(source=Path("/tmp/dump.rdb"), kind=InputSourceKind.LOCAL_RDB)],
        report_language="en-US",
    )
    monkeypatch.setattr("dba_assistant.capabilities.redis_rdb_analysis.service._parse_rdb_rows", lambda _path: rows)

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
    )

    assert isinstance(result, AnalysisReport)
    assert result.title == "Redis RDB Analysis Report"
    assert result.summary is not None
    assert result.summary.startswith("The analysis covers 1 sample, 3 keys, and 224 bytes.")
    assert "Expiration is configured for part of the dataset." in result.summary


def test_analyze_rdb_database_backed_route_stages_rows_into_shared_mysql_session(
    monkeypatch,
) -> None:
    rows = json.loads(Path("tests/fixtures/rdb/direct/sample_key_records.json").read_text(encoding="utf-8"))
    request = RdbAnalysisRequest(
        prompt="analyze this rdb via mysql",
        inputs=[SampleInput(source=Path("/tmp/dump.rdb"), kind=InputSourceKind.LOCAL_RDB)],
        path_mode="database_backed_analysis",
    )
    calls: dict[str, object] = {}

    monkeypatch.setattr(
        "dba_assistant.capabilities.redis_rdb_analysis.service._stream_rdb_rows",
        lambda _path: _streamed_rows(rows),
    )

    def fake_stage_rdb_rows_to_mysql(
        table_name: str,
        parsed_rows: list[dict[str, object]],
        *,
        source_file: str = "manual",
        run_id: str = "manual",
    ) -> dict[str, object]:
        calls["stage"] = (table_name, parsed_rows)
        return {"table": table_name, "staged": len(parsed_rows), "run_id": run_id}

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
        stage_rdb_rows_to_mysql=fake_stage_rdb_rows_to_mysql,
        analyze_staged_rdb_rows=_build_mysql_side_analysis(rows),
    )

    assert isinstance(result, AnalysisReport)
    staged_table, staged_rows = calls["stage"]
    assert staged_table.startswith("rdb_stage_auto_")
    assert staged_rows == rows
    assert result.metadata["route"] == "database_backed_analysis"
    assert result.metadata["path"] == "3a"
    assert result.metadata["mysql_full_table_reload"] == "disabled"
    assert any(section.id == "top_big_keys" for section in result.sections)


def test_analyze_rdb_database_backed_route_round_trips_mysql_text_values_without_type_breakage(
    monkeypatch,
) -> None:
    rows = [
        {
            "key_name": "cache:1",
            "key_type": "string",
            "size_bytes": 123,
            "has_expiration": False,
            "ttl_seconds": None,
        }
    ]
    request = RdbAnalysisRequest(
        prompt="analyze this rdb via mysql",
        inputs=[SampleInput(source=Path("/tmp/dump.rdb"), kind=InputSourceKind.LOCAL_RDB)],
        path_mode="database_backed_analysis",
    )
    adaptor = InMemoryTextMySQLAdaptor()
    config = MySQLConnectionConfig(
        host="localhost",
        port=3306,
        user="test",
        password="test",
        database="testdb",
    )

    monkeypatch.setattr(
        "dba_assistant.capabilities.redis_rdb_analysis.service._stream_rdb_rows",
        lambda _path: _streamed_rows(rows),
    )

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
        stage_rdb_rows_to_mysql=lambda table_name, parsed_rows, *, source_file="manual", run_id="manual": json.loads(
            stage_rdb_rows_to_mysql(
            adaptor,
            config,
            table_name,
            parsed_rows,
            run_id=run_id,
            source_file=source_file,
        )
        ),
        analyze_staged_rdb_rows=_build_mysql_side_analysis(rows),
    )

    assert isinstance(result, AnalysisReport)
    assert result.summary is not None
    assert result.summary.startswith("本次分析共覆盖 1 个样本、1 个键，累计内存占用 123 字节。")
    assert result.metadata["route"] == "database_backed_analysis"
    assert result.metadata["path"] == "3a"


def test_analyze_rdb_database_backed_route_uses_mysql_side_aggregation_without_full_reload(
    monkeypatch,
) -> None:
    rows = [
        {"key_name": "cache:1", "key_type": "string", "size_bytes": 300, "has_expiration": False, "ttl_seconds": None},
        {"key_name": "cache:2", "key_type": "hash", "size_bytes": 200, "has_expiration": True, "ttl_seconds": 60},
        {"key_name": "session:1", "key_type": "string", "size_bytes": 100, "has_expiration": False, "ttl_seconds": None},
    ]
    request = RdbAnalysisRequest(
        prompt="analyze this rdb via mysql",
        inputs=[SampleInput(source=Path("/tmp/dump.rdb"), kind=InputSourceKind.LOCAL_RDB)],
        path_mode="database_backed_analysis",
    )
    calls: dict[str, object] = {}

    monkeypatch.setattr(
        "dba_assistant.capabilities.redis_rdb_analysis.service._stream_rdb_rows",
        lambda _path: _streamed_rows(rows),
    )

    def fake_stage_rdb_rows_to_mysql(
        table_name: str,
        parsed_rows: list[dict[str, object]],
        *,
        source_file: str = "manual",
        run_id: str = "manual",
    ) -> dict[str, object]:
        calls["stage"] = (table_name, parsed_rows)
        return {"table": table_name, "staged": len(parsed_rows), "run_id": run_id}

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
        stage_rdb_rows_to_mysql=fake_stage_rdb_rows_to_mysql,
        analyze_staged_rdb_rows=_build_mysql_side_analysis(rows),
    )

    assert isinstance(result, AnalysisReport)
    assert calls["stage"][1] == rows
    assert result.metadata["route"] == "database_backed_analysis"
    assert result.metadata["mysql_full_table_reload"] == "disabled"
    assert result.summary is not None


def test_analyze_rdb_database_backed_route_uses_one_shared_staging_table_for_multiple_inputs(
    monkeypatch,
) -> None:
    rows_by_path = {
        Path("/tmp/a.rdb"): [
            {"key_name": "cache:1", "key_type": "string", "size_bytes": 300, "has_expiration": False, "ttl_seconds": None}
        ],
        Path("/tmp/b.rdb"): [
            {"key_name": "session:1", "key_type": "hash", "size_bytes": 200, "has_expiration": True, "ttl_seconds": 30}
        ],
        Path("/tmp/c.rdb"): [
            {"key_name": "order:1", "key_type": "string", "size_bytes": 100, "has_expiration": False, "ttl_seconds": None}
        ],
    }
    request = RdbAnalysisRequest(
        prompt="analyze these rdb files via mysql",
        inputs=[
            SampleInput(source=Path("/tmp/a.rdb"), kind=InputSourceKind.LOCAL_RDB),
            SampleInput(source=Path("/tmp/b.rdb"), kind=InputSourceKind.LOCAL_RDB),
            SampleInput(source=Path("/tmp/c.rdb"), kind=InputSourceKind.LOCAL_RDB),
        ],
        path_mode="database_backed_analysis",
        mysql_stage_batch_size=2,
    )
    staged_tables: list[str] = []

    monkeypatch.setattr(
        "dba_assistant.capabilities.redis_rdb_analysis.service._stream_rdb_rows",
        lambda path: _streamed_rows(rows_by_path[path]),
    )

    def fake_stage_rdb_rows_to_mysql(
        table_name: str,
        parsed_rows: list[dict[str, object]],
        *,
        source_file: str = "manual",
        run_id: str = "manual",
    ) -> dict[str, object]:
        staged_tables.append(table_name)
        return {"table": table_name, "staged": len(parsed_rows), "run_id": run_id}

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
        stage_rdb_rows_to_mysql=fake_stage_rdb_rows_to_mysql,
        analyze_staged_rdb_rows=_build_mysql_side_analysis(
            rows_by_path[Path("/tmp/a.rdb")]
            + rows_by_path[Path("/tmp/b.rdb")]
            + rows_by_path[Path("/tmp/c.rdb")]
        ),
    )

    assert isinstance(result, AnalysisReport)
    assert len(set(staged_tables)) == 1
    assert result.metadata["route"] == "database_backed_analysis"
    assert result.metadata["mysql_stage_batch_size"] == "2"


def test_analyze_rdb_database_backed_route_prefers_stream_rows_result_over_parse_rows_result(
    monkeypatch,
) -> None:
    rows = [
        {"key_name": "cache:1", "key_type": "string", "size_bytes": 300, "has_expiration": False, "ttl_seconds": None},
        {"key_name": "cache:2", "key_type": "hash", "size_bytes": 200, "has_expiration": True, "ttl_seconds": 60},
    ]
    request = RdbAnalysisRequest(
        prompt="analyze this rdb via mysql",
        inputs=[SampleInput(source=Path("/tmp/dump.rdb"), kind=InputSourceKind.LOCAL_RDB)],
        path_mode="database_backed_analysis",
    )

    class FakeStrategy:
        def parse_rows_result(self, _path: Path):
            raise AssertionError("database-backed analysis should not call parse_rows_result")

        def stream_rows_result(self, _path: Path):
            return StreamedRowsResult(rows=iter(rows), strategy_name="fake-stream-strategy")

    monkeypatch.setattr(
        "dba_assistant.capabilities.redis_rdb_analysis.service.build_default_rdb_parser_strategy",
        lambda: FakeStrategy(),
    )

    staged_rows: list[dict[str, object]] = []

    def fake_stage_rdb_rows_to_mysql(
        table_name: str,
        parsed_rows: list[dict[str, object]],
        *,
        source_file: str = "manual",
        run_id: str = "manual",
    ) -> dict[str, object]:
        del table_name, source_file, run_id
        staged_rows.extend(parsed_rows)
        return {"staged": len(parsed_rows), "run_id": "fake-run"}

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
        stage_rdb_rows_to_mysql=fake_stage_rdb_rows_to_mysql,
        analyze_staged_rdb_rows=_build_mysql_side_analysis(rows),
    )

    assert isinstance(result, AnalysisReport)
    assert staged_rows == rows
    assert result.metadata["parser_strategy"] == "fake-stream-strategy"


def test_analyze_rdb_large_local_file_auto_uses_streaming_large_summary_profile(
    monkeypatch,
    tmp_path: Path,
) -> None:
    large_rdb = tmp_path / "large.rdb"
    with large_rdb.open("wb") as handle:
        handle.truncate(1025 * 1024 * 1024)

    rows = [
        {"key_name": "cache:1", "key_type": "string", "size_bytes": 300, "has_expiration": False, "ttl_seconds": None},
        {"key_name": "loan:1", "key_type": "hash", "size_bytes": 200, "has_expiration": True, "ttl_seconds": 60},
        {"key_name": "queue:1", "key_type": "list", "size_bytes": 100, "has_expiration": False, "ttl_seconds": None},
    ]
    request = RdbAnalysisRequest(
        prompt="analyze this rdb",
        inputs=[SampleInput(source=large_rdb, kind=InputSourceKind.LOCAL_RDB)],
        path_mode="auto",
    )

    monkeypatch.setattr(
        "dba_assistant.capabilities.redis_rdb_analysis.service._stream_rdb_rows",
        lambda _path: _streamed_rows(rows),
    )

    class FailIfUsedCollector:
        def __init__(self, *args, **kwargs) -> None:
            raise AssertionError("large-file direct analysis should bypass PathCDirectParserCollector")

    monkeypatch.setattr(
        "dba_assistant.capabilities.redis_rdb_analysis.service.PathCDirectParserCollector",
        FailIfUsedCollector,
    )

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
    )

    assert isinstance(result, AnalysisReport)
    assert result.metadata["route"] == "direct_rdb_analysis"
    assert result.metadata["profile"] == "large_rdb_summary"
    assert result.metadata["large_rdb_protection"] == "enabled"
    assert all(
        section.id not in {"sample_overview", "top_string_keys", "top_hash_keys", "focused_prefix_analysis"}
        for section in result.sections
    )
    assert {section.id for section in result.sections} >= {
        "overall_summary",
        "key_type_summary",
        "expiration_summary",
        "prefix_top_summary",
        "top_big_keys",
    }


def test_large_rdb_summary_threshold_is_one_gibibyte() -> None:
    assert service_module.LARGE_RDB_AUTO_SUMMARY_THRESHOLD_BYTES == 1024 * 1024 * 1024


def test_analyze_rdb_remote_discovery_requires_rdb_path() -> None:
    request = RdbAnalysisRequest(
        prompt="analyze remote redis",
        inputs=[SampleInput(source="10.0.0.8:6379", kind=InputSourceKind.REMOTE_REDIS)],
    )

    with pytest.raises(ValueError, match="remote_discovery did not return rdb_path"):
        analyze_rdb(
            request,
            profile=None,
            remote_discovery=lambda *_args, **_kwargs: {},
        )


@pytest.mark.skipif(not HDT_BINARY.exists(), reason="HDT3213/rdb binary is not available in this workspace")
def test_analyze_rdb_uses_hdt_parser_strategy_for_v11_fixture() -> None:
    parser_strategy_module.build_default_rdb_parser_strategy.cache_clear()

    request = RdbAnalysisRequest(
        prompt="analyze a Redis 7 function dump",
        inputs=[
            SampleInput(
                source=Path("tests/fixtures/rdb/high_version/redis_v11_function.rdb"),
                kind=InputSourceKind.LOCAL_RDB,
            )
        ],
    )

    result = analyze_rdb(
        request,
        profile=None,
        remote_discovery=lambda *_args, **_kwargs: {"rdb_path": "/data/redis/dump.rdb"},
    )

    assert isinstance(result, AnalysisReport)
    assert result.metadata["route"] == "direct_rdb_analysis"
    assert result.metadata["path"] == "3c"
    assert result.metadata["parser_strategy"] == "HdtRdbCliStrategy"
    assert result.metadata["parser_binary"] == str(HDT_BINARY.resolve())
    assert result.summary is not None
    assert result.summary.startswith("本次分析共覆盖 1 个样本、0 个键，累计内存占用 0 字节。")

    parser_strategy_module.build_default_rdb_parser_strategy.cache_clear()
