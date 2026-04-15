from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document

from dba_assistant.core.analyzer.types import AnalysisResult, ReportSection, TableModel
from dba_assistant.core.reporter.docx_reporter import DocxReporter
from dba_assistant.core.reporter.report_model import AnalysisReport, ReportSectionModel, TableBlock, TextBlock, render_summary_text
from dba_assistant.core.reporter.types import OutputMode, ReportFormat, ReportOutputConfig


WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def build_analysis() -> AnalysisResult:
    return AnalysisResult(
        title="Redis RDB Analysis",
        summary="No urgent risk found.",
        sections=[
            ReportSection(
                title="Largest Keys",
                summary="Largest keys in the dataset",
                paragraphs=["Two session keys dominate memory usage."],
                tables=[
                    TableModel(
                        title="Top Keys",
                        columns=["Key", "Bytes"],
                        rows=[["session:1", "2048"]],
                    )
                ],
            )
        ],
        metadata={"environment": "prod", "generated_at": "2026-04-01"},
        risk_summary={"warning": 1},
    )


def test_docx_reporter_creates_a_minimal_report_document(tmp_path: Path) -> None:
    output_path = tmp_path / "rdb-analysis.docx"

    artifact = DocxReporter().render(
        build_analysis(),
        ReportOutputConfig(
            output_path=output_path,
            mode=OutputMode.REPORT,
            format=ReportFormat.DOCX,
            template_name="rdb-analysis",
            language="en-US",
        ),
    )

    document = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert artifact.output_path == output_path
    assert output_path.exists()
    assert "Redis RDB Analysis Report" in text
    assert "Largest Keys" in text
    assert "Risk Summary" not in text
    assert len(document.tables) == 1


def test_docx_reporter_supports_generic_analysis_report(tmp_path: Path) -> None:
    output_path = tmp_path / "generic-report.docx"
    report = AnalysisReport(
        title="Redis RDB 分析报告",
        summary="未发现紧急风险。",
        sections=[ReportSectionModel(id="summary", title="摘要", blocks=[TextBlock(text="ok")])],
        language="zh-CN",
    )

    artifact = DocxReporter().render(
        report,
        ReportOutputConfig(
            output_path=output_path,
            mode=OutputMode.REPORT,
            format=ReportFormat.DOCX,
            template_name="rdb-analysis",
            language="zh-CN",
        ),
    )

    document = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert artifact.output_path == output_path
    assert "Redis RDB 分析报告" in text
    assert "摘要" in text


def test_docx_reporter_applies_professional_styles_and_table_header_formatting(tmp_path: Path) -> None:
    output_path = tmp_path / "styled-report.docx"
    report = AnalysisReport(
        title="Redis RDB 分析报告",
        summary="本次分析覆盖 1 个样本，暂未发现额外确定性高风险。",
        sections=[
            ReportSectionModel(id="overview", title="样本与总体概况", level=1),
            ReportSectionModel(
                id="sample_overview",
                title="样本概览",
                level=2,
                blocks=[
                    TextBlock(text="本次分析纳入 1 个输入样本。"),
                    TableBlock(
                        title="样本清单",
                        columns=["样本名称", "样本类型", "数据来源"],
                        rows=[["sample-1", "本地 RDB", "/tmp/dump.rdb"]],
                    ),
                ],
            ),
        ],
        language="zh-CN",
    )

    DocxReporter().render(
        report,
        ReportOutputConfig(
            output_path=output_path,
            mode=OutputMode.REPORT,
            format=ReportFormat.DOCX,
            template_name="rdb-analysis",
            language="zh-CN",
        ),
    )

    styles_xml, document_xml = _read_docx_xml(output_path)
    document_root = ET.fromstring(document_xml)

    assert "DBA Cover Title" in styles_xml
    assert "DBA Body Text" in styles_xml
    assert 'w:eastAsia="宋体"' in styles_xml
    assert 'w:fill="D9E2F3"' in document_xml

    first_row = document_root.find(".//w:tbl/w:tr", WORD_NS)
    assert first_row is not None
    assert first_row.find("./w:trPr/w:tblHeader", WORD_NS) is None
    assert first_row.find(".//w:tcPr/w:shd", WORD_NS) is not None
    assert first_row.find(".//w:rPr/w:b", WORD_NS) is not None


def test_docx_reporter_numbers_headings_and_resets_minor_numbers_per_major_section(
    tmp_path: Path,
) -> None:
    output_path = tmp_path / "numbered-report.docx"
    report = AnalysisReport(
        title="Redis RDB 分析报告",
        summary="本次分析共覆盖 1 个样本，建议优先关注 string 类型高占用键。",
        sections=[
            ReportSectionModel(id="overview", title="样本与总体概况", level=1),
            ReportSectionModel(
                id="sample_overview",
                title="样本概览",
                level=2,
                blocks=[TextBlock(text="样本来源于本地 RDB 文件。")],
            ),
            ReportSectionModel(
                id="overall_summary",
                title="总体概览",
                level=2,
                blocks=[TextBlock(text="总体指标如下。")],
            ),
            ReportSectionModel(id="distribution_analysis", title="数据分布分析", level=1),
            ReportSectionModel(
                id="key_type_summary",
                title="键类型分布概览",
                level=2,
                blocks=[TextBlock(text="类型分布如下。")],
            ),
        ],
        language="zh-CN",
    )

    DocxReporter().render(
        report,
        ReportOutputConfig(
            output_path=output_path,
            mode=OutputMode.REPORT,
            format=ReportFormat.DOCX,
            template_name="rdb-analysis",
            language="zh-CN",
        ),
    )

    document = Document(output_path)
    non_empty_paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    texts = [paragraph.text for paragraph in non_empty_paragraphs]
    style_by_text = {paragraph.text: paragraph.style.name for paragraph in non_empty_paragraphs}

    assert texts.count("一、执行摘要") == 1
    assert "二、样本与总体概况" in texts
    assert "2.1 样本概览" in texts
    assert "2.2 总体概览" in texts
    assert "三、数据分布分析" in texts
    assert "3.1 键类型分布概览" in texts
    assert style_by_text["二、样本与总体概况"] == "DBA Heading 1"
    assert style_by_text["2.1 样本概览"] == "DBA Heading 2"


def test_docx_reporter_renders_one_word_table_per_table_block_without_repeat_header(
    tmp_path: Path,
) -> None:
    output_path = tmp_path / "single-table.docx"
    report = AnalysisReport(
        title="Redis RDB 分析报告",
        sections=[
            ReportSectionModel(id="big_key_analysis", title="大 Key 分析", level=1),
            ReportSectionModel(
                id="top_big_keys",
                title="总体大 Key 排名（Top 100）",
                level=2,
                blocks=[
                    TableBlock(
                        title="总体大 Key 排名（Top 100）",
                        columns=["键名", "键类型", "内存占用（字节）"],
                        rows=[[f"key:{index}", "string", str(index)] for index in range(1, 101)],
                    )
                ],
            ),
        ],
        language="zh-CN",
    )

    DocxReporter().render(
        report,
        ReportOutputConfig(
            output_path=output_path,
            mode=OutputMode.REPORT,
            format=ReportFormat.DOCX,
            template_name="rdb-analysis",
            language="zh-CN",
        ),
    )

    document = Document(output_path)
    _, document_xml = _read_docx_xml(output_path)

    assert len(document.tables) == 1
    assert document_xml.count("<w:tbl>") == 1
    assert document_xml.count("键名") == 1
    assert "w:tblHeader" not in document_xml


def test_docx_reporter_preserves_focus_only_report_scope_without_full_sections(
    tmp_path: Path,
) -> None:
    output_path = tmp_path / "focus-only.docx"
    report = AnalysisReport(
        title="Redis RDB 分析报告",
        summary="本报告仅输出用户指定的重点前缀详情。",
        sections=[
            ReportSectionModel(id="focused_prefix_analysis", title="重点前缀详情分析", level=1),
            ReportSectionModel(
                id="focused_prefix_detail:tag:*",
                title="前缀 tag:* 详情",
                level=2,
                blocks=[
                    TextBlock(text="此前缀范围共匹配 2 个键。"),
                    TableBlock(
                        title="前缀 tag:* Top Keys（Top 10）",
                        columns=["键名", "键类型", "内存占用（字节）"],
                        rows=[["tag:1", "string", "500"]],
                    ),
                ],
            ),
        ],
        metadata={"profile": "rcs", "scope": "focused_prefix_only"},
        language="zh-CN",
    )

    summary_text = render_summary_text(report, language="zh-CN")
    assert "样本与总体概况" not in summary_text
    assert "重点前缀详情分析" in summary_text

    DocxReporter().render(
        report,
        ReportOutputConfig(
            output_path=output_path,
            mode=OutputMode.REPORT,
            format=ReportFormat.DOCX,
            template_name="rdb-analysis",
            language="zh-CN",
        ),
    )

    document = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "重点前缀详情分析" in text
    assert "前缀 tag:* 详情" in text
    assert "样本与总体概况" not in text


def test_docx_reporter_numbers_sub_tables_with_three_level_prefix(tmp_path: Path) -> None:
    """Round 1.2 #1: Table titles within level-2 sections get {major}.{minor}.{sub} numbering."""
    output_path = tmp_path / "sub-numbered.docx"
    report = AnalysisReport(
        title="Redis 巡检报告",
        sections=[
            ReportSectionModel(id="redis_db", title="Redis 数据库检查", level=1),
            ReportSectionModel(
                id="redis_db__cluster-a",
                title="cluster-a",
                level=2,
                blocks=[
                    TextBlock(text="检查详情"),
                    TableBlock(title="架构与角色摘要", columns=["节点", "角色"], rows=[["n1", "master"]]),
                    TableBlock(title="版本与一致性检查", columns=["版本", "节点数"], rows=[["7.0", "1"]]),
                ],
            ),
            ReportSectionModel(
                id="redis_db__cluster-b",
                title="cluster-b",
                level=2,
                blocks=[
                    TextBlock(text="检查详情"),
                    TableBlock(title="架构与角色摘要", columns=["节点", "角色"], rows=[["n2", "master"]]),
                ],
            ),
        ],
        language="zh-CN",
    )

    DocxReporter().render(
        report,
        ReportOutputConfig(
            output_path=output_path,
            mode=OutputMode.REPORT,
            format=ReportFormat.DOCX,
            template_name="rdb-analysis",
            language="zh-CN",
        ),
    )

    document = Document(output_path)
    texts = [p.text for p in document.paragraphs if p.text.strip()]

    # Major section: 一、Redis 数据库检查
    assert any("一、Redis 数据库检查" in t for t in texts)
    # Minor heading: 1.1 cluster-a
    assert "1.1 cluster-a" in texts
    # Sub table titles within cluster-a: 1.1.1, 1.1.2
    assert any("1.1.1 架构与角色摘要" in t for t in texts)
    assert any("1.1.2 版本与一致性检查" in t for t in texts)
    # Minor heading: 1.2 cluster-b
    assert "1.2 cluster-b" in texts
    # Sub table title within cluster-b resets: 1.2.1
    assert any("1.2.1 架构与角色摘要" in t for t in texts)


def test_docx_reporter_anomaly_highlighting_applies_red_to_keywords(tmp_path: Path) -> None:
    """Round 1.2 #5: Anomaly keywords should get red+bold formatting in table cells."""
    output_path = tmp_path / "highlighted.docx"
    report = AnalysisReport(
        title="Redis 巡检报告",
        sections=[
            ReportSectionModel(
                id="risk",
                title="风险总结",
                level=1,
                blocks=[
                    TableBlock(
                        title="风险清单",
                        columns=["问题", "等级", "说明"],
                        rows=[
                            ["OOM command not allowed", "critical", "内存不足导致异常"],
                            ["Normal operation", "info", "一切正常"],
                        ],
                    ),
                ],
            ),
        ],
        language="zh-CN",
    )

    DocxReporter().render(
        report,
        ReportOutputConfig(
            output_path=output_path,
            mode=OutputMode.REPORT,
            format=ReportFormat.DOCX,
            template_name="rdb-analysis",
            language="zh-CN",
        ),
    )

    _, document_xml = _read_docx_xml(output_path)
    # Verify red color (FF0000) appears in the document for anomaly keywords
    assert "FF0000" in document_xml, "Anomaly keywords should be highlighted in red"


def test_default_output_path_for_inspection_defaults_to_configured_artifact_policy(monkeypatch) -> None:
    from dba_assistant.core.reporter.output_path_policy import DEFAULT_ARTIFACT_DIR, default_report_output_path

    monkeypatch.setattr(
        "dba_assistant.core.reporter.output_path_policy._timestamp_slug",
        lambda: "20260414_120000",
    )
    path = default_report_output_path("docx", report_slug="inspection")
    assert path.parent == DEFAULT_ARTIFACT_DIR
    assert "dba_assistant_redis_inspection" in str(path)


def _read_docx_xml(output_path: Path) -> tuple[str, str]:
    with ZipFile(output_path) as archive:
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        document_xml = archive.read("word/document.xml").decode("utf-8")
    return styles_xml, document_xml
